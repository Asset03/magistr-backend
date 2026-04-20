import logging
from typing import Dict, List, Any, Optional, Union
import os
import re
from pathlib import Path
import json

logger = logging.getLogger(__name__)

class DocumentParser:
    def __init__(self):
        """Initialize document parser with support for multiple file formats"""
        try:
            self.supported_formats = {
                '.txt': self._parse_txt,
                '.json': self._parse_json,
                '.csv': self._parse_csv,
                '.xml': self._parse_xml,
                '.html': self._parse_html
            }
            
            # Try to import optional parsers
            self.pdf_parser = None
            self.docx_parser = None
            
            try:
                import pdfplumber
                self.pdf_parser = pdfplumber
                logger.info("PDF parser (pdfplumber) loaded")
            except ImportError:
                logger.warning("PDF parser not available. Install pdfplumber for PDF support.")
            
            try:
                import docx
                self.docx_parser = docx
                logger.info("DOCX parser (python-docx) loaded")
            except ImportError:
                logger.warning("DOCX parser not available. Install python-docx for DOCX support.")
            
            logger.info("Document parser initialized successfully")
            
        except Exception as e:
            logger.error(f"Error initializing document parser: {str(e)}")
            raise

    def parse_document(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Parse document and extract text content"""
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                return {'error': f'File not found: {file_path}'}
            
            # Get file extension
            ext = file_path.suffix.lower()
            
            # Check if format is supported
            if ext not in self.supported_formats and ext not in ['.pdf', '.docx']:
                return {'error': f'Unsupported file format: {ext}'}
            
            # Parse based on format
            if ext == '.pdf':
                return self._parse_pdf(file_path, **kwargs)
            elif ext == '.docx':
                return self._parse_docx(file_path, **kwargs)
            else:
                parser_func = self.supported_formats[ext]
                return parser_func(file_path, **kwargs)
                
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {str(e)}")
            return {'error': str(e)}

    def _parse_txt(self, file_path: Path, **kwargs) -> Dict[str, Any]:
        """Parse plain text file"""
        try:
            encoding = kwargs.get('encoding', 'utf-8')
            
            with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
                content = file.read()
            
            return {
                'text': content,
                'metadata': {
                    'file_name': file_path.name,
                    'file_size': file_path.stat().st_size,
                    'format': 'txt',
                    'encoding': encoding
                },
                'success': True
            }
            
        except Exception as e:
            logger.error(f"Error parsing TXT file: {str(e)}")
            return {'error': str(e)}

    def _parse_json(self, file_path: Path, **kwargs) -> Dict[str, Any]:
        """Parse JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Extract text from JSON structure
            text_content = self._extract_text_from_json(data)
            
            return {
                'text': text_content,
                'metadata': {
                    'file_name': file_path.name,
                    'file_size': file_path.stat().st_size,
                    'format': 'json',
                    'structure_keys': list(data.keys()) if isinstance(data, dict) else []
                },
                'success': True
            }
            
        except Exception as e:
            logger.error(f"Error parsing JSON file: {str(e)}")
            return {'error': str(e)}

    def _parse_csv(self, file_path: Path, **kwargs) -> Dict[str, Any]:
        """Parse CSV file"""
        try:
            delimiter = kwargs.get('delimiter', ',')
            quote_char = kwargs.get('quote_char', '"')
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                lines = file.readlines()
            
            # Simple CSV parsing
            rows = []
            for line in lines:
                # Remove quotes and split by delimiter
                row = line.strip().replace(quote_char, '').split(delimiter)
                rows.append(row)
            
            # Convert to text
            text_content = '\n'.join([' '.join(row) for row in rows])
            
            return {
                'text': text_content,
                'metadata': {
                    'file_name': file_path.name,
                    'file_size': file_path.stat().st_size,
                    'format': 'csv',
                    'rows': len(rows),
                    'columns': len(rows[0]) if rows else 0
                },
                'success': True
            }
            
        except Exception as e:
            logger.error(f"Error parsing CSV file: {str(e)}")
            return {'error': str(e)}

    def _parse_xml(self, file_path: Path, **kwargs) -> Dict[str, Any]:
        """Parse XML file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            # Simple XML text extraction (remove tags)
            text_content = re.sub(r'<[^>]+>', ' ', content)
            text_content = re.sub(r'\s+', ' ', text_content).strip()
            
            return {
                'text': text_content,
                'metadata': {
                    'file_name': file_path.name,
                    'file_size': file_path.stat().st_size,
                    'format': 'xml'
                },
                'success': True
            }
            
        except Exception as e:
            logger.error(f"Error parsing XML file: {str(e)}")
            return {'error': str(e)}

    def _parse_html(self, file_path: Path, **kwargs) -> Dict[str, Any]:
        """Parse HTML file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            # Remove script and style tags
            content = re.sub(r'<script[^>]*>.*?</script>', ' ', content, flags=re.DOTALL | re.IGNORECASE)
            content = re.sub(r'<style[^>]*>.*?</style>', ' ', content, flags=re.DOTALL | re.IGNORECASE)
            
            # Remove HTML tags
            text_content = re.sub(r'<[^>]+>', ' ', content)
            text_content = re.sub(r'\s+', ' ', text_content).strip()
            
            return {
                'text': text_content,
                'metadata': {
                    'file_name': file_path.name,
                    'file_size': file_path.stat().st_size,
                    'format': 'html'
                },
                'success': True
            }
            
        except Exception as e:
            logger.error(f"Error parsing HTML file: {str(e)}")
            return {'error': str(e)}

    def _parse_pdf(self, file_path: Path, **kwargs) -> Dict[str, Any]:
        """Parse PDF file using pdfplumber"""
        try:
            if not self.pdf_parser:
                return {'error': 'PDF parser not available. Install pdfplumber.'}
            
            text_content = []
            metadata = {
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'format': 'pdf',
                'pages': 0
            }
            
            with self.pdf_parser.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(f"--- Page {page_num} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num}: {str(e)}")
                        continue
            
            return {
                'text': '\n\n'.join(text_content),
                'metadata': metadata,
                'success': True
            }
            
        except Exception as e:
            logger.error(f"Error parsing PDF file: {str(e)}")
            return {'error': str(e)}

    def _parse_docx(self, file_path: Path, **kwargs) -> Dict[str, Any]:
        """Parse DOCX file using python-docx"""
        try:
            if not self.docx_parser:
                return {'error': 'DOCX parser not available. Install python-docx.'}
            
            doc = self.docx_parser.Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(' | '.join(row_data))
                tables_text.append('\n'.join(table_data))
            
            # Combine all text
            text_content = '\n\n'.join(paragraphs)
            if tables_text:
                text_content += '\n\n--- Tables ---\n\n' + '\n\n'.join(tables_text)
            
            metadata = {
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'format': 'docx',
                'paragraphs': len(paragraphs),
                'tables': len(doc.tables)
            }
            
            return {
                'text': text_content,
                'metadata': metadata,
                'success': True
            }
            
        except Exception as e:
            logger.error(f"Error parsing DOCX file: {str(e)}")
            return {'error': str(e)}

    def _extract_text_from_json(self, data: Any, max_depth: int = 10) -> str:
        """Recursively extract text from JSON structure"""
        try:
            if max_depth <= 0:
                return ''
            
            if isinstance(data, str):
                return data
            elif isinstance(data, (int, float, bool)):
                return str(data)
            elif isinstance(data, list):
                return ' '.join(self._extract_text_from_json(item, max_depth - 1) for item in data)
            elif isinstance(data, dict):
                return ' '.join(
                    f"{key}: {self._extract_text_from_json(value, max_depth - 1)}"
                    for key, value in data.items()
                )
            else:
                return str(data)
                
        except Exception as e:
            logger.warning(f"Error extracting text from JSON: {str(e)}")
            return ''

    def batch_parse_documents(self, file_paths: List[str], **kwargs) -> List[Dict[str, Any]]:
        """Parse multiple documents"""
        results = []
        
        for file_path in file_paths:
            try:
                result = self.parse_document(file_path, **kwargs)
                result['file_path'] = file_path
                results.append(result)
            except Exception as e:
                logger.error(f"Error parsing {file_path}: {str(e)}")
                results.append({
                    'file_path': file_path,
                    'error': str(e),
                    'success': False
                })
        
        return results

    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        formats = list(self.supported_formats.keys())
        
        if self.pdf_parser:
            formats.append('.pdf')
        
        if self.docx_parser:
            formats.append('.docx')
        
        return formats
